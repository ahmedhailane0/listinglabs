"""Build PERPS_REPORT.docx — polished Word document combining EXPLANATION.txt
content with the analysis.ipynb charts and per-chart explanations.

Workflow:
  1. Re-execute analysis.ipynb if analysis_run.ipynb is missing/stale.
  2. Extract every PNG output from the executed notebook into charts/.
  3. Parse EXPLANATION.txt into sections.
  4. Assemble a styled docx: cover -> sections -> charts -> closing.
"""
import base64
import json
import re
import subprocess
import sys
from datetime import datetime
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

HERE = Path(__file__).parent
NB = HERE / "analysis.ipynb"
NB_RUN = HERE / "analysis_run.ipynb"
EXPLAIN = HERE / "EXPLANATION.txt"
CHARTS = HERE.parent / "charts"
OUT = HERE / "PERPS_REPORT.docx"
CSV_PATH = HERE / "enriched_clean.csv"

# -------------------- design tokens --------------------

NAVY     = RGBColor(0x1F, 0x3A, 0x5F)
NAVY_HEX = "1F3A5F"
GOLD     = RGBColor(0xC8, 0xA0, 0x4C)
GOLD_HEX = "C8A04C"
TEXT     = RGBColor(0x30, 0x30, 0x30)
MUTED    = RGBColor(0x77, 0x77, 0x77)
ACCENT   = RGBColor(0xC0, 0x39, 0x2B)  # for warnings / emphasis

CALLOUT_BG  = "EAF2FA"   # soft blue
PANEL_BG    = "F5F5F5"   # light gray
HEAD_BAR_BG = "1F3A5F"   # navy band

# -------------------- pipeline plumbing --------------------

def ensure_executed():
    if NB_RUN.exists() and NB_RUN.stat().st_mtime > NB.stat().st_mtime:
        return
    print("re-executing notebook ...")
    subprocess.run(
        [sys.executable, "-m", "jupyter", "nbconvert", "--to", "notebook",
         "--execute", str(NB), "--output", str(NB_RUN.name),
         "--ExecutePreprocessor.timeout=180"],
        check=True, cwd=HERE,
    )


def extract_charts():
    CHARTS.mkdir(exist_ok=True)
    for f in CHARTS.glob("*.png"):
        f.unlink()
    nb = json.loads(NB_RUN.read_text(encoding="utf-8"))
    chart_files = []
    n = 0
    for cell in nb["cells"]:
        if cell.get("cell_type") != "code":
            continue
        for out in cell.get("outputs", []):
            data = out.get("data") or {}
            if "image/png" not in data:
                continue
            n += 1
            png_path = CHARTS / f"chart_{n}.png"
            png_path.write_bytes(base64.b64decode(data["image/png"]))
            chart_files.append(png_path)
    print(f"extracted {len(chart_files)} chart PNGs to {CHARTS}")
    return chart_files


def parse_explanation():
    text = EXPLAIN.read_text(encoding="utf-8")
    parts = re.split(r"^=+\s*$", text, flags=re.MULTILINE)
    sections = []
    i = 1
    while i < len(parts):
        heading = parts[i].strip()
        body = parts[i + 1].strip() if i + 1 < len(parts) else ""
        if heading and body and not heading.lower().startswith("end of explanation"):
            sections.append((heading, body))
        i += 2
    return sections


def dataset_stats():
    """Quick stats for the cover-page card."""
    df = pd.read_csv(CSV_PATH, parse_dates=["binance_perp_date"])
    return {
        "tokens": len(df),
        "date_min": df["binance_perp_date"].min().strftime("%b %Y"),
        "date_max": df["binance_perp_date"].max().strftime("%b %Y"),
        "columns": len(df.columns),
    }


# -------------------- XML / docx helpers --------------------

def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def set_cell_border(cell, *, left=None, right=None, top=None, bottom=None):
    """Set per-side borders on a table cell. Each value: dict(sz, color)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for side, spec in (("left", left), ("right", right),
                       ("top", top), ("bottom", bottom)):
        el = OxmlElement(f"w:{side}")
        if spec is None:
            el.set(qn("w:val"), "nil")
        else:
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), str(spec.get("sz", 4)))
            el.set(qn("w:color"), spec.get("color", "000000"))
        tc_borders.append(el)


def add_horizontal_rule(paragraph, color_hex="1F3A5F", size=12):
    """Append a horizontal rule (paragraph bottom border) to a paragraph."""
    p_pr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pbdr.append(bottom)
    p_pr.append(pbdr)


def add_page_number_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def add_header(doc, text):
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED
    run.italic = True


def add_h1(doc, text):
    """Major section heading: navy, large, with gold underline."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = NAVY
    run.font.name = "Calibri"
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    add_horizontal_rule(p, color_hex=GOLD_HEX, size=12)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_body_paragraph(doc, text, size=11, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = TEXT
    if italic:
        run.italic = True
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.25
    return p


def add_mono_block(doc, text):
    """Indented monospace block — for data rows and code."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = TEXT


def add_callout(doc, title, bullets, bg_hex=CALLOUT_BG, border_color_hex=NAVY_HEX):
    """Single-cell colored card with title + bullet list."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = True
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, bg_hex)
    set_cell_border(
        cell,
        left={"sz": 24, "color": border_color_hex},
        top={"sz": 4, "color": "D0D0D0"},
        right={"sz": 4, "color": "D0D0D0"},
        bottom={"sz": 4, "color": "D0D0D0"},
    )
    # title
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = NAVY
    # bullets
    for b in bullets:
        bp = cell.add_paragraph()
        bp.paragraph_format.left_indent = Inches(0.1)
        bp.paragraph_format.space_after = Pt(3)
        r = bp.add_run("•  " + b)
        r.font.size = Pt(10.5)
        r.font.color.rgb = TEXT
    doc.add_paragraph().paragraph_format.space_after = Pt(2)  # spacer


def add_stats_card(doc, stats):
    """Cover-page stats panel: 4 cells side by side."""
    items = [
        ("Tokens analyzed", f"{stats['tokens']}"),
        ("Time window", f"{stats['date_min']} — {stats['date_max']}"),
        ("Data sources", "Binance · OKX · Coinbase · Bithumb · Upbit · CryptoRank · RootData"),
        ("Columns per token", f"{stats['columns']}"),
    ]
    tbl = doc.add_table(rows=2, cols=len(items))
    tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for col, (label, value) in enumerate(items):
        # top cell: label
        c1 = tbl.cell(0, col)
        set_cell_bg(c1, NAVY_HEX)
        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p1.add_run(label.upper())
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # bottom cell: value
        c2 = tbl.cell(1, col)
        set_cell_bg(c2, "FFFFFF")
        set_cell_border(c2, left={"sz": 4, "color": "DDDDDD"},
                        right={"sz": 4, "color": "DDDDDD"},
                        bottom={"sz": 4, "color": "DDDDDD"})
        p2 = c2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p2.add_run(value)
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = NAVY


def add_chart_card(doc, idx, chart_file, title, what, how, takeaway):
    """Each chart: colored title bar, image, then a 3-row info table."""
    doc.add_page_break()
    # navy title bar
    bar = doc.add_table(rows=1, cols=1)
    cell = bar.cell(0, 0)
    set_cell_bg(cell, NAVY_HEX)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # image
    pic_p = doc.add_paragraph()
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_p.add_run().add_picture(str(chart_file), width=Inches(6.3))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_r = cap.add_run(f"Figure {idx} — {title.split('—', 1)[-1].strip()}")
    cap_r.italic = True
    cap_r.font.size = Pt(9)
    cap_r.font.color.rgb = MUTED
    cap.paragraph_format.space_after = Pt(12)

    # 3-row explanation table
    info = doc.add_table(rows=3, cols=2)
    info.autofit = True
    rows_data = [
        ("WHAT YOU'RE LOOKING AT", what),
        ("HOW TO READ IT", how),
        ("WHAT IT TELLS US", takeaway),
    ]
    for i, (label, text) in enumerate(rows_data):
        c_label = info.cell(i, 0)
        c_text = info.cell(i, 1)
        c_label.width = Inches(1.6)
        c_text.width = Inches(4.7)
        set_cell_bg(c_label, NAVY_HEX)
        set_cell_bg(c_text, "FAFAFA" if i % 2 == 0 else "FFFFFF")
        set_cell_border(c_label, top={"sz": 4, "color": "DDDDDD"},
                        bottom={"sz": 4, "color": "DDDDDD"})
        set_cell_border(c_text, top={"sz": 4, "color": "DDDDDD"},
                        bottom={"sz": 4, "color": "DDDDDD"},
                        right={"sz": 4, "color": "DDDDDD"})
        # label
        pl = c_label.paragraphs[0]
        pl.paragraph_format.space_before = Pt(6)
        pl.paragraph_format.space_after = Pt(6)
        rl = pl.add_run(label)
        rl.bold = True
        rl.font.size = Pt(9)
        rl.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # text
        pt = c_text.paragraphs[0]
        pt.paragraph_format.space_before = Pt(6)
        pt.paragraph_format.space_after = Pt(6)
        pt.paragraph_format.left_indent = Inches(0.1)
        rt = pt.add_run(text)
        rt.font.size = Pt(10.5)
        rt.font.color.rgb = TEXT


# -------------------- body-text router --------------------

KEY_PATTERN = re.compile(
    r"(ρ\s*=\s*[+−-]?\d+\.\d+|"
    r"rho\s*=\s*[+−-]?\d+\.\d+|"
    r"R²\s*=\s*\d+\.\d+|"
    r"p\s*<\s*0\.\d+|"
    r"\$[\d,\.]+(?:K|M|B)?|"
    r"[+−-]\d+\.\d+%?)"
)


def add_styled_paragraph(doc, line):
    """Body line with inline highlights: bold + color on key metrics."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    parts = KEY_PATTERN.split(line)
    for j, part in enumerate(parts):
        run = p.add_run(part)
        run.font.size = Pt(11)
        run.font.color.rgb = TEXT
        if j % 2 == 1:
            run.bold = True
            run.font.color.rgb = NAVY


def add_body_text(doc, body):
    """Render section body. Recognises:
       - 'Example N —' blocks  -> sub-heading + monospace data block
       - lines starting with '---' surrounding example -> stripped
       - lines beginning with '|' or aligned with whitespace + '|' -> mono table
       - bullet markers ( '-', '•', '*', or 'a)' style) -> bullet list
       - plain prose paragraphs -> styled body with inline highlights
    """
    for block in re.split(r"\n\s*\n", body):
        if not block.strip():
            continue
        lines = [ln for ln in block.splitlines()
                 if not re.fullmatch(r"-{5,}", ln.strip())]
        if not lines:
            continue
        clean = "\n".join(lines).strip()
        if not clean:
            continue

        # Worked example heading
        m = re.match(r"^(Example \d+\s*[—-].+)$", lines[0].strip())
        if m:
            # spacer
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
            hp = doc.add_paragraph()
            hp.paragraph_format.space_before = Pt(8)
            hp.paragraph_format.space_after = Pt(2)
            run = hp.add_run(m.group(1))
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = NAVY
            # gold underline
            add_horizontal_rule(hp, color_hex=GOLD_HEX, size=6)
            body_text = "\n".join(lines[1:]).strip()
            if body_text:
                add_mono_block(doc, body_text)
            continue

        # ASCII table / data
        if any("|" in line and re.match(r"^\s*\S", line) for line in lines) \
                or re.search(r"^\s*\w+\s+\|", clean, re.M):
            add_mono_block(doc, clean)
            continue

        # Bullet-list block (lines starting with - or • or a))
        if all(re.match(r"^\s*[-•*]\s|^\s*[a-z]\)\s|^\s*\d+\.\s", ln)
               or not ln.strip() for ln in lines):
            for ln in lines:
                if not ln.strip():
                    continue
                bullet_text = re.sub(r"^\s*[-•*]\s*", "", ln)
                bullet_text = re.sub(r"^\s*([a-z]\)|\d+\.)\s*", r"\1 ", bullet_text)
                bp = doc.add_paragraph()
                bp.paragraph_format.left_indent = Inches(0.25)
                bp.paragraph_format.space_after = Pt(3)
                bp.paragraph_format.line_spacing = 1.2
                bp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                r = bp.add_run("•  " + bullet_text.strip())
                r.font.size = Pt(11)
                r.font.color.rgb = TEXT
            continue

        # Default: prose paragraph(s) — collapse soft line breaks
        # If lines look like wrapped prose (no leading whitespace),
        # join them with spaces; otherwise emit per-line.
        looks_prose = all(not re.match(r"^\s{2,}", ln) for ln in lines)
        if looks_prose:
            add_styled_paragraph(doc, " ".join(ln.strip() for ln in lines))
        else:
            for ln in lines:
                if not ln.strip():
                    continue
                add_styled_paragraph(doc, ln)


# -------------------- chart descriptions --------------------

CHART_DESCRIPTIONS = [
    (
        "Chart 1 — The 2×3 distributions grid",
        "Six small charts in one image, each showing the distribution of one key column across all 344 tokens. Top row: exchange-footprint and post-launch returns. Bottom row: volume decay, FDV at launch, and VC funding size.",
        "Bar height (or histogram height) = how many tokens fall in each bucket. The return histograms are clipped to ±300% so a few extreme outliers don't squash the rest. FDV and funding panels use a log-10 scale because token sizes span many orders of magnitude.",
        "Returns and volume decay are long-tailed and slightly negative on average — most perps drift down after launch. Money columns are heavy-tailed (a few giants, many small ones). This is exactly why we use Spearman ranks rather than raw values for the correlation math.",
    ),
    (
        "Chart 2 — CryptoRank ICO raised vs RootData total funding",
        "A scatter plot comparing the two funding data sources. Each dot is a project that has a value in BOTH sources. X-axis is CryptoRank's public-sale (ICO/IDO) raised amount; Y-axis is RootData's total funding (all rounds combined). Both axes are log scale. Dashed diagonal = perfect agreement.",
        "Read the gap from the diagonal: dots above the line mean RootData reports MORE money than the CryptoRank ICO figure — expected, since RootData includes private rounds. The chart should look mostly empty — sparsity itself is the finding.",
        "Only 4 projects had both numbers, because CryptoRank's free API only exposes public-sale data and most modern projects skip public sales. The two sources can't be compared on USD raised. We cross-check them instead on investor counts (ρ ≈ 0.58 — reasonable agreement) and treat RootData as authoritative for funding totals.",
    ),
    (
        "Chart 3 — Spearman correlation matrix (the headline)",
        "A grid of 10 'before' features (rows) crossed with 7 'after' outcomes (columns). Each cell shows ρ between that pair, computed across tokens with non-null values for both. Cell colors: deep blue = strong inverse, white = no relationship, deep red = strong positive.",
        "Hunt for the deepest reds and blues — those are the strongest signals. For example, row n_exchanges_before × column vol_decay_ratio is the biggest red square (ρ ≈ +0.22): tokens already on more venues retain volume better. The rd_total_funding_usd row is mostly pale — funding USD doesn't predict outcomes.",
        "Three patterns dominate. (1) Prior exchange footprint matters for VOLUME RETENTION, not for price. (2) FDV at launch matters for PRICE RETURN — bigger holds up better. (3) Co-launches hurt both volume AND returns. These are bivariate; section 6 uses regression to disentangle overlapping effects.",
    ),
    (
        "Chart 4 — Returns by exchange-footprint bucket",
        "Two side-by-side boxplots. X-axis groups tokens by how many of the six tracked venues had listed them before the Binance perp. Left panel = 7-day return; right panel = 30-day return. Box anatomy: orange line = median, box = middle 50%, whiskers = the rest, circles = outliers.",
        "If more prior listings meant better price performance, the median line would climb left-to-right. It mostly doesn't — medians wobble in a narrow band. Widest boxes are buckets with fewer tokens (high n_exchanges_before is rare), where individual outliers stretch the box.",
        "Prior exchange footprint does NOT translate to stronger price returns over 7 or 30 days. The heatmap signal for this feature was about VOLUME sustain, not price. This chart is the 'absence-of-effect' chart for price — a useful negative result.",
    ),
    (
        "Chart 5 — Returns and volume decay by funding bucket",
        "Two side-by-side boxplots grouped by funding size (per RootData): under $1M, $1–10M, $10–100M, over $100M. Left = 30-day return per bucket; right = volume decay ratio per bucket.",
        "If well-funded projects outperformed, the median would rise from left to right. In both panels it doesn't move much. The largest bucket (>$100M) usually has few tokens and a wide box; the smallest bucket likewise — large boxes there are sample-size noise.",
        "Funding amount alone is not predictive of price or volume at the perp horizon. The relevant funding signal is the COUNT of investors (more investors → faster volume fade), not the dollar AMOUNT raised.",
    ),
]


# -------------------- main build --------------------

def build_report(chart_files):
    doc = Document()
    # tighter margins for a roomier feel
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = TEXT

    add_header(doc, "Binance Perps · Correlation Analysis")
    add_page_number_footer(doc)
    stats = dataset_stats()

    # ===================== Cover page =====================
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(60)

    eyebrow = doc.add_paragraph()
    eyebrow.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = eyebrow.add_run("RESEARCH REPORT")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = GOLD
    r.font.name = "Calibri"

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Binance Perps")
    r.bold = True
    r.font.size = Pt(38)
    r.font.color.rgb = NAVY

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Correlation analysis of launches, listings, funding and volume")
    r.font.size = Pt(15)
    r.font.color.rgb = MUTED
    r.italic = True

    # decorative rule
    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_horizontal_rule(rule, color_hex=GOLD_HEX, size=12)

    # date / author
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run(datetime.now().strftime("%B %Y"))
    r.font.size = Pt(11)
    r.font.color.rgb = MUTED

    for _ in range(3):
        doc.add_paragraph()

    add_stats_card(doc, stats)

    for _ in range(2):
        doc.add_paragraph()

    add_callout(
        doc,
        "Bottom line",
        [
            "Tokens listed on more exchanges BEFORE the Binance perp keep their "
            "trading volume better post-launch  (ρ = +0.22, p < 0.001).",
            "Tokens co-launched on many venues within ±7 days fade FASTER on both "
            "volume and 30-day price  (ρ between −0.14 and −0.17). Classic "
            "sell-the-news.",
            "Higher FDV at launch predicts better 30-day returns  (ρ = +0.19, "
            "p < 0.001). Big tokens hold up better.",
            "More VC investors correlates with faster volume fade  (ρ = −0.16). "
            "Hype at launch, profit-taking after.",
            "Total VC funding dollars are NOT predictive of price or volume "
            "outcomes. The investor COUNT matters, the dollar AMOUNT does not.",
        ],
    )

    doc.add_page_break()

    # ===================== Contents =====================
    add_h1(doc, "Contents")
    sections = parse_explanation()
    for i, (sec_heading, _) in enumerate(sections, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"  {i}.  {sec_heading}")
        r.font.size = Pt(11)
        r.font.color.rgb = TEXT
    extras = [
        ("THE CHARTS, ONE AT A TIME", "5 figures with what / how / takeaway"),
        ("Closing thoughts", "summary + caveats + what's next"),
    ]
    for title_extra, desc in extras:
        p = doc.add_paragraph()
        r = p.add_run(f"  •  {title_extra}  ")
        r.font.size = Pt(11)
        r.font.color.rgb = TEXT
        r.bold = True
        r2 = p.add_run(f"— {desc}")
        r2.font.size = Pt(10.5)
        r2.italic = True
        r2.font.color.rgb = MUTED

    doc.add_page_break()

    # ===================== Main body =====================
    for sec_heading, body in sections:
        add_h1(doc, sec_heading)
        add_body_text(doc, body)

    # ===================== Charts =====================
    doc.add_page_break()
    add_h1(doc, "The charts, one at a time")
    add_body_paragraph(
        doc,
        "Below are the five figures produced by analysis.ipynb on enriched_clean.csv. "
        "Each is followed by a three-row table — what you're looking at, how to read "
        "it, and what it tells us — so the chart can stand alone without the notebook.",
        italic=True,
    )

    for i, (title, what, how, takeaway) in enumerate(CHART_DESCRIPTIONS):
        if i >= len(chart_files):
            break
        add_chart_card(doc, i + 1, chart_files[i], title, what, how, takeaway)

    # ===================== Closing =====================
    doc.add_page_break()
    add_h1(doc, "Closing thoughts")
    add_body_paragraph(
        doc,
        "The headline finding is the volume-retention one: tokens that already had a "
        "footprint on other CEXes before the Binance perp launch keep their trading "
        "volume better in the weeks afterward. Co-launches (multiple venues lighting "
        "up the same week) do the opposite — they pile up day-1 volume that then "
        "fades fast.",
    )
    add_body_paragraph(
        doc,
        "For price, the strongest predictor is FDV at launch: bigger tokens "
        "outperform over 30 days. Funding background (count of investors) is "
        "associated with faster volume fade, but the dollar amount of funding "
        "doesn't predict anything cleanly.",
    )

    add_callout(
        doc,
        "Caveats and next steps",
        [
            "These are bivariate Spearman correlations; effects overlap with each "
            "other and with FDV. The OLS in section 6 of the notebook only starts "
            "to disentangle them — it is not a rigorous causal model.",
            "Pull 4-hour returns instead of daily to catch more of the immediate "
            "post-launch action.",
            "Add open-interest growth data — arguably a cleaner 'is this perp "
            "catching on' signal than raw volume.",
            "Add sector tags (meme / DeFi / AI / L1 / gaming) so patterns can be "
            "compared by category.",
        ],
        bg_hex=PANEL_BG,
        border_color_hex=GOLD_HEX,
    )

    doc.save(OUT)
    print(f"wrote {OUT}")


def main():
    ensure_executed()
    chart_files = extract_charts()
    if len(chart_files) < 5:
        print(f"WARNING: only {len(chart_files)} charts found, expected 5.")
    build_report(chart_files)


if __name__ == "__main__":
    main()
